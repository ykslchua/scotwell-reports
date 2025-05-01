from docx import Document
from utils import replace_in_paragraphs, add_photos_to_service_report
from flask import request
import os
from docx.shared import Inches, Pt

def generate_service_report(form_data, photo_paths_before=None, photo_paths_after=None):
    doc = Document('templates/service_report_template.docx')  # Ensure doc is initialized

    # Define your placeholder replacements
    placeholders = {
        "Company Name:": f"{form_data['company_name']}",
        "Site Address:": f"{form_data['site_address']}",
        "Service Date:": f"{form_data['service_date']}",
        "Screen Condition:": f"{form_data['screen_condition']}",
        "Action Taken:": f"{form_data['action_taken']}",
        "Follow Up / Recommendations:": f"{form_data['follow_up']}",
        "Engineer Name:": f"Service Engineer: {form_data['technician_name']}",
        "Chargeable Spare Parts:": f"Chargeable Spare Parts: {form_data['chargeable_spare_parts']}",
        "Chargeable Module Repair:": f"Chargeable Module Repair: {form_data['chargeable_module_repair']}"
    }

    # Replace in document paragraphs
    replace_in_paragraphs(doc.paragraphs, placeholders)

    # Replace in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs, placeholders)

    # Find "Parts Repair/Replacement" section
    insert_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if "Parts Repair/Replacement" in paragraph.text:
            insert_index = i
            break

    # Create spare parts table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    headers = ['Spare Part', 'Model', 'Quantity', 'Replaced', 'Repaired']
    hdr_cells = table.rows[0].cells

    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True

    # Adjust column widths
    column_widths = [Inches(2), Inches(3), Inches(1), Inches(0.75), Inches(0.75)]
    for i, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = column_widths[i]

    # Populate table rows
    spare_parts = ['LED Module', 'Power Supply', 'Receiving Card', 'Hub Card', 'Sending Card', 'Data Cable', 'Others']

    spare_part_name = request.form.getlist('spare_part_name[]')
    spare_part_model = request.form.getlist('spare_part_model[]')
    spare_part_qty = request.form.getlist('spare_part_qty[]')
    spare_part_replaced = request.form.getlist('spare_part_replaced[]')
    spare_part_repaired = request.form.getlist('spare_part_repaired[]')

    submitted_parts = {}
    for i in range(len(spare_part_name)):
        name = spare_part_name[i]
        model = spare_part_model[i] if i < len(spare_part_model) else ''
        qty = spare_part_qty[i] if i < len(spare_part_qty) else ''
        is_replaced = "☑" if str(i) in spare_part_replaced else "☐"
        is_repaired = "☑" if str(i) in spare_part_repaired else "☐"
        submitted_parts[name] = {
            'model': model,
            'qty': qty,
            'replaced': is_replaced,
            'repaired': is_repaired
        }

    for part in spare_parts:
        model = submitted_parts.get(part, {}).get('model', '')
        qty = submitted_parts.get(part, {}).get('qty', '')
        replaced = submitted_parts.get(part, {}).get('replaced', '☐')
        repaired = submitted_parts.get(part, {}).get('repaired', '☐')

        row_cells = table.add_row().cells
        row_cells[0].text = part
        row_cells[1].text = model
        row_cells[2].text = qty
        row_cells[3].text = replaced
        row_cells[4].text = repaired

    # Insert table into document
    if insert_index is not None:
        paragraph = doc.paragraphs[insert_index]
        paragraph._element.addnext(table._element)
    else:
        doc.add_paragraph("Parts Repair/Replacement")
        doc._body._element.append(table._element)

    # Insert page break before "Before Photos"
    doc.add_paragraph().add_run().add_break()

    # Add Before Photos
    add_photos_to_service_report(doc, photo_paths_before, 'Before Photos')  # Pass doc as argument

    # Add a blank paragraph (line break) before After Photos
    doc.add_paragraph()
    add_photos_to_service_report(doc, photo_paths_after, 'After Photos')  # Pass doc as argument

    # Save the document
    output_filename = f"{form_data['company_name']}_{form_data['site_address']}_{form_data.get('service_date', '')}_report.docx"
    output_path = os.path.join('static', 'reports', output_filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    return output_filename
