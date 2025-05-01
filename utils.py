import os
import pandas as pd
import openai
import configparser
import io
import piexif
from docx.shared import Inches, Pt
from flask import request
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont
from datetime import datetime


# -------------------------------
# File-related utilities
# -------------------------------

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def get_image_timestamp(image_file):
    try:
        filename = getattr(image_file, 'filename', '').lower()

        # JPEG or JPG
        if filename.endswith('.jpg') or filename.endswith('.jpeg'):
            image_file.seek(0)
            img = Image.open(image_file)
            if 'exif' in img.info:
                exif_dict = piexif.load(img.info['exif'])
                date_str = exif_dict['0th'].get(piexif.ImageIFD.DateTime)
                if date_str:
                    return datetime.strptime(date_str.decode(), "%Y:%m:%d %H:%M:%S").strftime("%Y-%m-%d %H:%M:%S")

    except Exception as e:
        print(f"[get_image_timestamp] Metadata extraction failed: {e}")

    return None


def sanitize_folder_name(company_name):
    return company_name.replace(';', '_').replace('/', '_').replace('\\', '_')


# -------------------------------
# Dropdown Data Loader
# -------------------------------

def load_dropdown_data(file_path='data/dropdownlist.xlsx'):
    df = pd.read_excel(file_path)

    company_names = df['Company Name'].dropna().unique().tolist()
    site_addresses_by_company = df.groupby('Company Name')['Site Address'].apply(list).to_dict()

    chargeable_spare_parts = df.drop_duplicates(subset=['Site Address']) \
        .set_index('Site Address')['Chargeable Spare Parts'].to_dict()

    chargeable_module_repair = df.drop_duplicates(subset=['Site Address']) \
        .set_index('Site Address')['Chargeable Module Repair'].to_dict()

    return company_names, site_addresses_by_company, chargeable_spare_parts, chargeable_module_repair


# -------------------------------
# OpenAI Helper
# -------------------------------
# Load OpenAI API key
config = configparser.ConfigParser()
config.read("config.ini")
openai.api_key = config.get("OPENAI", "API_KEY", fallback=None)

if not openai.api_key:
    print("Error: OpenAI API key is missing. Set it in 'config.ini' under [OPENAI] section.")
    exit(1)
    
def elaborate_description(field_name, description):
    if not description.strip():
        return ""

    if field_name.lower() == "screen condition":
        prompt = f"""
        You are an LED billboard maintenance expert. Reword the following screen condition issue into a clear, concise, and proper sentence that describes the problem without adding extra details:

        Issue: {description}

        Ensure the response is technical, precise, and limited to the problem observed.
        """
    else:
        prompt = f"""
        You are an expert in LED billboard maintenance. Provide a detailed and professional elaboration for any action_taken:

        Action: {description}

        Ensure the response is technical, concise, and actionable.
        """

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert in LED billboard maintenance. You will receive short field notes and must convert them into clear, professional, and technically precise sentences."},
                {"role": "user", "content": description}
            ]
        )
        return response["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"Error calling OpenAI API: {e}")
        return description


def make_sentence(prompt_text):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful assistant who rewrites short technical notes into complete, professional sentences."},
            {"role": "user", "content": f"Turn the following into a proper sentence:\n{prompt_text}"}
        ],
        temperature=0.4,
        max_tokens=100
    )
    return response.choices[0].message.content.strip()

# -------------------------------
# Document Helpers
# -------------------------------

def replace_in_paragraphs(paragraphs, placeholders):
    for para in paragraphs:
        for placeholder, replacement in placeholders.items():
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, replacement)

# Function to add photos in rows of 2
def add_photos_to_service_report(doc, photo_paths, title):
    if photo_paths:
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(14)

        # Create a table with two columns per row
        num_rows = (len(photo_paths) + 1) // 2  # Calculate how many rows we need
        table = doc.add_table(rows=num_rows, cols=2)
        table.style = 'Table Grid'  # Optional, you can adjust the style

        for i, path in enumerate(photo_paths):
            row_index = i // 2  # Determine the row based on index
            col_index = i % 2   # Determine the column (0 or 1)

            # Ensure that the cell exists
            row_cells = table.rows[row_index].cells

            # Add the image and filename
            picture_cell = row_cells[col_index].paragraphs[0].add_run()
            picture_cell.add_picture(path, width=Inches(2.5))  # Adjust the width as needed
            
            

def add_photos_to_visual_report(doc, photo_paths, title, insert_into_cell=None):
    if not photo_paths or insert_into_cell is None:
        return

    # Add a section title (e.g. "Before Photos") with spacing
    if title:
        p_title = insert_into_cell.add_paragraph()
        p_title.add_run(title).bold = True
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after = Pt(2)

    # Insert each photo with spacing
    for photo_path in photo_paths:
        p = insert_into_cell.add_paragraph()
        run = p.add_run()
        run.add_picture(photo_path, width=Inches(2))
        p.paragraph_format.space_after = Pt(6)



def add_timestamp_to_image(image_file, timestamp=None):
    timestamp = timestamp or get_image_timestamp(image_file)
    image_file.seek(0)

    if not timestamp:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    img = Image.open(image_file).convert("RGB")
    width, height = img.size
    draw = ImageDraw.Draw(img)

    dynamic_font_size = max(80, min(40, height // 25))
    try:
        font = ImageFont.truetype("arial.ttf", size=dynamic_font_size)
    except:
        font = ImageFont.load_default()

    margin = 10
    raise_amount = int(height * 0.05)
    position = (margin, height - raise_amount - dynamic_font_size)

    try:
        text_bbox = draw.textbbox((0, 0), timestamp, font=font)
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]
    except AttributeError:
        text_width, text_height = draw.textsize(timestamp, font=font)

    rect_x0, rect_y0 = position
    rect_x1, rect_y1 = rect_x0 + text_width + 6, rect_y0 + text_height + 4
    draw.rectangle([rect_x0, rect_y0, rect_x1, rect_y1], fill="black")
    draw.text((position[0] + 3, position[1] + 2), timestamp, fill="white", font=font)

    output = io.BytesIO()
    img.save(output, format="JPEG")
    output.seek(0)
    return output




