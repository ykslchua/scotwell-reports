from docx import Document
from flask import request
from utils import replace_in_paragraphs, add_photos_to_visual_report
import os

def generate_visual_report(form_data, photo_paths_before=None, photo_paths_after=None):
    doc = Document('templates/visual_report_template.docx')  # Use a visual report template


    # Helper to return checkbox-style string
    def checkbox_line(label, value):
        return f"{label}:  {'☑' if value == 'Yes' else '☐'} Yes   {'☐' if value == 'Yes' else '☑'} No"


    placeholders = {
        "Company Name:": form_data['company_name'],
        "Site Address:": form_data['site_address'],
        "Inspection Date:": form_data['service_date'],
        "Reporting Engineer:": f"Reporting Engineer: {form_data['technician_name']}",
        "Visual Inspection:": checkbox_line("Visual Inspection", form_data.get('visual_inspection', 'Yes')),
        "Faults Reported:": checkbox_line("Faults Reported", form_data.get('faults_reported', 'No')),
        "Work Description:": form_data['work_description']
    }

    # Replace placeholders in document
    replace_in_paragraphs(doc.paragraphs, placeholders)


    # Replace placeholders and find the "Screen condition:" cell
    target_cell = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs, placeholders)
                if "Screen condition:" in cell.text:
                    target_cell = cell  # Save for later photo insertion

    # Only show "Before Photos" header if faults are reported
    if photo_paths_before:
        before_title = "Before Photos" if form_data.get('faults_reported') == 'Yes' else None
        add_photos_to_visual_report(doc, photo_paths_before, before_title, insert_into_cell=target_cell)

    # Add After Photos only if faults are reported
    if form_data.get('faults_reported') == 'Yes' and photo_paths_after:
        add_photos_to_visual_report(doc, photo_paths_after, "After Photos", insert_into_cell=target_cell)


    # Save report
    filename = f"{form_data['company_name']}_{form_data['site_address']}_{form_data['service_date']}_visual_report.docx"
    output_path = os.path.join('static', 'reports', filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    return filename
